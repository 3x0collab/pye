import os
import datetime


import csv
import json
import openpyxl
import xml.etree.ElementTree as ET

from sqlalchemy import create_engine,inspect
from sqlalchemy.orm import sessionmaker
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.declarative import declarative_base
from django.core.exceptions import ObjectDoesNotExist
from customer.models import TaskConnectorData,FileModels,Connections
from utils.vortex_pdf_parser import VortexPdfParser
from utils.vortex_ingester import VortexIngester
from utils.pye_bot import WebScraperBot
from .email import Mail
import docx


class Connectors(object):
	"""docstring for Connectors"""
	def __init__(self, connector=None,task=None):
		super(Connectors, self).__init__()
		self.task = task
		self.connector = connector
		self.credentials = json.loads(connector.parameters)
		self.name = connector.name
		self.connector_type = connector.connector_type
		self.data = []
		print(f"Connecting to {self.name} Connector")
		self.connect()
	
	def connect(self):
		pass

	def get(self):
		return self.data

	def post(self,data):
		self.data = data

	def log_connection(self,method="GET",data=[]):
		log(self.name+f"_{self.connector.pk}_{method}_logs",f"""Connector {method} method run succesfully, with {len(data)} records.\n""")
		



class DatasetConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(DatasetConnector, self).__init__(connector,task) 
		

	def connect(self):
		# print('here now', self.credentials)
		for x in range(int(self.credentials['datasets_number'])):
			self.data.append(json.loads(self.credentials['datasets_content']))
		self.log_connection('GET',self.data)

	def post(self,data):
		log(self.name+"-"+str(self.connector.pk),data)
		self.log_connection('POST',data)


class RawConnector(Connectors):
	"""docstring for RawConnector"""
	def __init__(self,connector=None,task=None):
		super(RawConnector, self).__init__(connector,task) 

	def connect(self):
		self.data = self.credentials['data']
		self.log_connection('GET',self.data)

	def post(self,data):
		log(self.name+"-"+str(self.connector.pk),data)
		self.log_connection('POST',data)


class LinksConnector(Connectors):
	"""docstring for RawConnector"""
	def __init__(self,connector=None,task=None):
		super(LinksConnector, self).__init__(connector,task) 

	def connect(self):
		if self.task.name != 'Test Task':
			bot = WebScraperBot(self.credentials.get('url'), self.credentials.get('depth'), self.credentials.get('read_documents_flag'))
			result = bot.run_bot(self.credentials.get('depth'),self.credentials.get('follow_external'), self.credentials.get('follow_internal'))
			text_content,external_links, internal_links = result 
			self.data = text_content
			self.log_connection('GET',external_links + internal_links)

	def post(self,data):
		log(self.name+"-"+str(self.connector.pk),data)
		self.log_connection('POST',data)




class FileConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(FileConnector, self).__init__(connector,task)

	def connect(self):
		if self.task.name != 'Test Task':
			save_data = []
			print(FileModels.objects.filter(connector_id=self.connector.pk).exclude(ingest=1))
			for file_model in  FileModels.objects.filter(connector_id=self.connector.pk).exclude(ingest=1):
				file_name = file_model.file.name
				data = self.read_file(
					file_model.filetype,
					file_model.file,
					self.credentials.get("delimiter","|"),
					self.credentials.get("quotechar","'")
					)
				# print(data)
				if type(data) == list:
					save_data = save_data +data
				else:
					save_data.append(data)
			self.data = save_data
			self.log_connection('GET',self.data)
 
		

	def read_file(self, file_type, file, delimiter=None, quotechar=None):
		file_path = file.path
		file_type = file_type.lower()
		if file_type == 'csv':
			with open(file_path, 'r') as f:
				reader = csv.DictReader(f, delimiter=delimiter, quotechar=quotechar)
				data = [dict(row) for row in reader]
			return data
		elif file_type == 'txt':
		    with open(file_path, 'r') as f:
		        data = f.readlines()
		    return data
		elif file_type == 'excel':
			workbook = openpyxl.load_workbook(filename=file_path)
			sheet = workbook.active
			data = []
			headers = [cell.value for cell in sheet[1]]
			for row in sheet.iter_rows(min_row=2):
				row_data = {headers[i]: cell.value for i, cell in enumerate(row)}
				data.append(row_data)
			return data
		elif file_type == 'json':
			with open(file_path, 'r') as f:
				data = json.load(f)
			return data
		elif file_type == 'xml':
			tree = ET.parse(file_path)
			root = tree.getroot()
			data = []
			for child in root:
				row_data = {subchild.tag: subchild.text for subchild in child}
				data.append(row_data)
			return data
		elif file_type == 'pdf':
			vortex_pdf_parser = VortexPdfParser()
			vortex_pdf_parser.set_pdf_file_path(file_path)
			document_chunks = vortex_pdf_parser.clean_text_to_docs()
			return document_chunks
		elif file_type == 'docx' or file_type == 'doc':
			doc = docx.Document(file_path)
			content = '\n'.join([p.text for p in doc.paragraphs])
			return content
		else:
			raise ValueError(f'Unsupported file type: {file_type}') 



	def post(self,data):
		log(self.name+"-"+str(self.connector.pk),data)
		self.log_connection('POST',data)

 



class DatabaseConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(DatabaseConnector, self).__init__(connector,task)

	def connect(self):
		self.connect_to_database()
		if self.connector.pk != 'test':
			self.data = self.pull_data()
			self.log_connection('GET',self.data)



	def connect_to_database(self):

		# Extract form inputs
		db_type = str(self.credentials['db_type']).lower()
		host = self.credentials['host']
		port = self.credentials['port']
		db_name = self.credentials['db_name']
		username = self.credentials['username']
		password = self.credentials['password']
		sql_query = self.credentials['sql_query']

		# Database connection URL
		if db_type == 'postgresql':
			db_url = f'postgresql://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'mysql':
			db_url = f'mysql+pymysql://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'sqlite':
			db_url = f'sqlite:///{db_name}'
		elif db_type == 'oracle':
			db_url = f'oracle://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'mssql':
			driver = 'ODBC Driver 17 for SQL Server'
			db_url = f'mssql+pyodbc://{username}:{password}@{host}:{port}/{db_name}?driver={driver}'
		elif db_type == 'amazon_redshift':
			db_url = f'redshift+psycopg2://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'google_bigquery':
			db_url = f'bigquery://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'snowflake':
			db_url = f'snowflake://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'apache_cassandra':
			db_url = f'cassandra://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'apache_hadoop':
			db_url = f'hadoop://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'teradata':
			db_url = f'teradata://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'ibm_db2':
			db_url = f'db2://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'sqlite':
			db_url = f'sqlite:///{db_name}'
		elif db_type == 'informix':
			db_url = f'informix://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'sap_hana':
			db_url = f'hana://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'vertica':
			db_url = f'vertica+pyodbc://{username}:{password}@{host}:{port}/{db_name}'
		elif db_type == 'MariaDB':
			db_url = 'mysql+pymysql://{username}:{password}@{host}:{port}/{db_name}'
		else:
			raise ValueError('Invalid database type')

		# Connect to the database 
		self.engine = create_engine(db_url)
		Session = sessionmaker(bind=self.engine)
		self.session = Session()

	def pull_data(self):
		# Execute SQL query
		table_name = self.credentials.get('table_name',"")
		sql_query = self.credentials['sql_query']
		query_params = self.get_query_params()

		if sql_query:
			result = self.session.execute(sql_query,query_params)
		else:
			result = self.session.execute(f'SELECT * FROM {table_name}' )

		column_names = [desc[0] for desc in result.cursor.description]  # Retrieve column names from the ResultSet
		data = []
		while True:
			rows = result.fetchmany(50)  # Retrieve 50 rows at a time
			if not rows:
				break
			data.extend([dict(zip(column_names, row)) for row in rows])

		return data

	def get_primary_key(self, table_name):
		Base = declarative_base(bind=self.engine)
		Base.metadata.reflect()

		# Check if the table exists in the reflected metadata
		if table_name.lower() in Base.metadata.tables.keys():
			# Get the table object
			table = Base.metadata.tables[table_name.lower()]
			inspector = inspect(self.engine)
			primary_key_columns = inspector.get_pk_constraint(table_name)['constrained_columns']
			primary_key_column = primary_key_columns[0] if primary_key_columns else None
			# Print the primary key column name
			if primary_key_column:
				if table_name.isupper():
					return primary_key_column.upper()
			return primary_key_column
		else:
			print("Table does not exist")

	
	def get_query_params(self):
		params = {}
		params['LAST_RUN_DATE'] = str(self.task.last_run_date)
		params['LAST_RUN_TIME'] = str(self.task.last_run_time)
		return params

	def post(self,data=[]):
		table_name = self.credentials.get('table_name',"")
		pk = self.get_primary_key(table_name)
		query_params = self.get_query_params()

		# Perform data updates or inserts
		if isinstance(data, list):
			for entry in data:
				try:
					self.session.execute(f"INSERT INTO {table_name} ({', '.join(entry.keys())}) VALUES ({', '.join([f':{key}' for key in entry.keys()])})", {**entry,**query_params})
				except IntegrityError:
					update_values = ', '.join([f"{key} = :{key}" for key in entry.keys()])
					self.session.execute(f"UPDATE {table_name} SET {update_values} WHERE {pk} = :{pk}", {**entry,**query_params})
		elif isinstance(data, dict):
			for key, value in data.items():
				for entry in value:
					try:
						self.session.execute(f"INSERT INTO {table_name} ({', '.join(entry.keys())}) VALUES ({', '.join([f':{key}' for key in entry.keys()])})", {**entry,**query_params})
					except IntegrityError:
						update_values = ', '.join([f"{key} = :{key}" for key in entry.keys()])
						self.session.execute(f"UPDATE {table_name} SET {update_values} WHERE {pk} = :{pk}", {**entry,**query_params})

		# Commit the changes
		self.session.commit()
		# Close the session
		self.session.close()
		if self.connector.pk != 'test':
			self.log_connection('POST',data)


	
class TaskConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(TaskConnector, self).__init__(connector,task)

	def connect(self):
		try:
			task_id = f"{self.credentials['tasks']}-{self.task.pk}"
			task_data = TaskConnectorData.objects.get(task_id=task_id )
			self.data = json.loads(task_data.data)
			task_data.delete()
		except ObjectDoesNotExist:
			self.log_connection('GET',f"No Data for Task: {self.task.pk}-{self.task.name} in Task Connector")

	def post(self,data=[]):
		try:
			task_id = f"{self.credentials['tasks']}-{self.task.pk}"
			task_data = TaskConnectorData.objects.get(task_id=task_id )
			task_data.data = json.dumps(data)
			task_data.save()
		except ObjectDoesNotExist:
			json_data = json.dumps(data)
			task_data = TaskConnectorData.objects.create(**{"task_id":task_id,"data":json_data}  )

			self.log_connection('POST',data)	



class EmailConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(EmailConnector, self).__init__(connector,task)


	def connect(self):
		try:
			self.mail_client = Mail(self.credentials['email'], self.credentials['password'], self.credentials['smtp_server'], self.credentials['smtp_port'])
			self.data = self.mail_client.read_emails(num_emails=50) 
			self.log_connection('GET',self.data)
		except ObjectDoesNotExist:
			self.log_connection('GET',f"No Data for Email: {self.task.pk}-{self.task.name} in Email Connector")


	def post(self,data=[]):
		# 0 - 'recipient' | 1 - 'subject'  |  2 - 'body'
		self.mail_client.send_email(data[0] ,data[1], data[2])
		self.log_connection('POST',data)
	


class TrainConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(TrainConnector, self).__init__(connector,task)

	def connect(self):
		pass

	def post(self,data=None):
		if data and len(data):
			try:
				ex_connections = Connections.objects.get(name = self.connector.name)
				pipeline_exist = ex_connections.pipelines.filter(name=self.task.name).count()
				if not pipeline_exist:
					ex_connections.pipelines.add(self.task)

			except Exception as e:
				create_connections = Connections.objects.create(
				name=self.connector.name,
				description=self.connector.description,
				connector = self.connector
					)
				create_connections.pipelines.add(self.task)
				create_connections.save()

			ingester = VortexIngester(str(data),f"{self.connector.name}_{str(self.connector.created_by)}")
			ingester.ingest()
			# print('dfd1',FileModels.objects.filter(connector_id=self.connector.pk).exclude(ingest=1))
			FileModels.objects.filter(connector_id=self.connector.pk).exclude(ingest=1).update(ingest=1)
			self.log_connection(self.connector.name,data)


class APIConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(APIConnector, self).__init__(connector,task)

	def connect(self):
		pass


class DAConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(APIConnector, self).__init__(connector,task)

	def connect(self):
		pass


class ApplicationConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(ApplicationConnector, self).__init__(connector,task)

	def connect(self):
		pass



class SocialConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(SocialConnector, self).__init__(connector,task)

	def connect(self):
		pass





class CloudConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(CloudConnector, self).__init__(connector,task)

	def connect(self):
		pass





class IOTConnector(Connectors):
	"""docstring for DatasetConnector"""
	def __init__(self,connector=None,task=None):
		super(IOTConnector, self).__init__(connector,task)

	def connect(self):
		pass




connector_map = {
	"database":DatabaseConnector,
	"flatfile":FileConnector,
	"api":APIConnector,
	"application":ApplicationConnector,
	"social_media":SocialConnector,
	"cloud":CloudConnector,
	"iot":IOTConnector,

}	

def pick_connector(Connector,task):

	if Connector.type_name == 'datasets':
		return DatasetConnector(Connector,task) 
	elif Connector.type_name == 'blank':
		return RawConnector(Connector,task) 
	elif Connector.type_name == 'pipeline':
		return TaskConnector(Connector,task) 
	elif Connector.type_name == 'pye da':
		return DAConnector(Connector,task) 
	elif Connector.type_name == 'ai model':
		return TrainConnector(Connector,task)
	elif Connector.type_name == 'email':
		return EmailConnector(Connector,task) 
	elif Connector.type_name == 'links':
		return LinksConnector(Connector,task) 
	else:
		con =  connector_map[Connector.connector_type](Connector,task) 
		return con


def log(name,data=[]):
		path = os.getcwd()
		name = name.replace(":","_")
		path = os.path.join(path,'logs',name)
		print('path',path)
		os.makedirs(path, exist_ok=True)
		date_strf = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
		log_file = os.path.join(path, f'{date_strf}.txt')

		with open(log_file, "a") as f:
			if isinstance(data, str):
				f.write(data)
			elif isinstance(data, (list, tuple,set)):
				for item in data:
					f.write(str(item) + "\n")
			elif isinstance(data, dict):
				for key, value in data.items():
					f.write(str(key) + ": " + str(value) + "\n")
			else:
				f.write(str(data))

				
# credentials = {"csrfmiddlewaretoken": "HMvy80qcSzVZpqrSANusg6k75MZQWwrIio9xFfDosyFAl69qMuPxvVVN1leoz1F4", "connector_type": "dummy", "connector_pk": "4", "datasets": "BasicNewsArticles", "datasets_content": "{\"article_id\":\"ART001\",\"headline\":\"COVID-19 Cases Continue to Rise Across the Globe\",\"article_text\":\"As the Omicron variant spreads, COVID-19 cases are surging worldwide. Health officials are urging people to get vaccinated and follow safety guidelines to help slow the spread of the virus.\",\"publish_date\":\"03/13/2023\",\"source\":\"CNN\",\"keywords\":[\"COVID-19\",\"Omicron variant\",\"health officials\",\"vaccinated\",\"safety guidelines\",\"slow the spread\"],\"entities\":[{\"entity_name\":\"COVID-19\",\"entity_type\":\"Disease\"},{\"entity_name\":\"Omicron\",\"entity_type\":\"Variant\"},{\"entity_name\":\"CNN\",\"entity_type\":\"Organization\"}]}", "datasets_number": "10"}
# dummy1 = DummyConnector(credentials)
# print(dummy1.name)
# dummy_data = dummy1.get()[0]
# print(dummy_data['article_id'])
	